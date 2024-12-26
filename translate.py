import os, sys, time
from openai import OpenAI
from docx import Document
from pypdf import PdfReader

# Export the OpenAI API key as OS env var
# Replace with your API key
os.environ['OPENAI_API_KEY']="REPLACE_WITH_YOUR_API_KEY"
client = OpenAI()


def translate_pdf_to_word(pdf_file_path, target_lang="Korean", output_file="translated_doc.docx"):
    """
    Translates a PDF file to English and a specified target language and saves it as a Word document.

    Args:
        pdf_file_path (str): The path to the PDF file.
        target_lang (str): The target language for translation (default: Korean).
        output_file (str): The name of the output Word file (default: translated_doc.docx).
    """

    # Extract text from PDF
    reader = PdfReader(pdf_file_path)
    text = ""

    # Create a Word document and write the translated text
    doc = Document()

    print(f"There are {len(reader.pages)} pages to translate.")
    for i, page in enumerate(reader.pages):
        print(f"Translating page {i+1}...")
        text += page.extract_text()

        # Translate text using OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
              {"role": "system", "content": f"You are a helpful assistant that translates text to English and {target_lang}. Place '### English Translation:' before the English-translated text. Then place '### {target_lang} translation:' before {target_lang}-translated text. The text in single quotes should be copied verbatim."},
              {"role": "user", "content": f"Translate the following text:\n\n{text}"}
            ]
        )
        translated_text = response.choices[0].message.content
        #print(translated_text)
        doc.add_paragraph("### Original text:")
        doc.add_paragraph(text)
        doc.add_paragraph(translated_text)

        # Sleep 5 sec to work around OpenAI's rate limiting
        doc.save(output_file)
        time.sleep(5)


if __name__ == "__main__":
    pdf_file = "input.pdf"
    if len(sys.argv) > 1:    # If provided, use the first argument as input file path
      pdf_file = sys.argv[1]

    translate_pdf_to_word(pdf_file, "Korean")
    print("Translation was successful. Check out the file: translated_doc.docx")
